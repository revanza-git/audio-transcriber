from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import os
import tempfile
import uuid
from pathlib import Path
from docx import Document
import aiofiles
from typing import Optional
import whisper
import torch
import asyncio
import signal
from contextlib import asynccontextmanager
import concurrent.futures
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Environment configuration
BACKEND_HOST = os.getenv("BACKEND_HOST", "0.0.0.0")
BACKEND_PORT = int(os.getenv("BACKEND_PORT", "8000"))
BACKEND_DEBUG = os.getenv("BACKEND_DEBUG", "False").lower() == "true"

# CORS configuration
CORS_ORIGINS = os.getenv("CORS_ALLOWED_ORIGINS", "http://localhost:5173,http://localhost:3000").split(",")
CORS_ORIGINS = [origin.strip() for origin in CORS_ORIGINS if origin.strip()]

# Timeout configuration
DEFAULT_TIMEOUT = int(os.getenv("DEFAULT_TIMEOUT_MINUTES", "10"))
MIN_TIMEOUT = int(os.getenv("MIN_TIMEOUT_MINUTES", "1"))
MAX_TIMEOUT = int(os.getenv("MAX_TIMEOUT_MINUTES", "60"))

# Directory configuration
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "uploads"))
OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "outputs"))

# Whisper model configuration
PREFERRED_MODEL = os.getenv("PREFERRED_WHISPER_MODEL", "large-v3")
FALLBACK_MODELS = os.getenv("FALLBACK_WHISPER_MODELS", "medium,base,tiny").split(",")
FALLBACK_MODELS = [model.strip() for model in FALLBACK_MODELS if model.strip()]

# Initialize FastAPI app
app = FastAPI(title="Record to Text API", version="1.0.0", debug=BACKEND_DEBUG)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load Whisper model
print("Record to Text API starting...")
print(f"Environment: {'Development' if BACKEND_DEBUG else 'Production'}")
print(f"CORS Origins: {CORS_ORIGINS}")
print("Loading Whisper model...")

def load_best_whisper_model():
    """Load the best Whisper model that the system can handle"""
    # Try preferred model first, then fallbacks
    models_to_try = [(PREFERRED_MODEL, "Preferred model")]
    for model in FALLBACK_MODELS:
        if model != PREFERRED_MODEL:
            models_to_try.append((model, "Fallback model"))
    
    model_descriptions = {
        "large-v3": "Best accuracy, requires more RAM/GPU",
        "large-v2": "Excellent accuracy, requires more RAM/GPU", 
        "large": "Great accuracy, requires more RAM/GPU",
        "medium": "Good balance of speed and accuracy",
        "small": "Faster, good accuracy",
        "base": "Fast and lightweight",
        "tiny": "Fastest, basic accuracy"
    }
    
    for model_name, category in models_to_try:
        try:
            description = model_descriptions.get(model_name, "Unknown model")
            print(f"Trying {model_name} model... ({category}: {description})")
            model = whisper.load_model(model_name)
            print(f"✅ Whisper {model_name} model loaded successfully!")
            return model, model_name, True
        except Exception as e:
            print(f"❌ Failed to load {model_name} model: {e}")
            continue
    
    print("❌ Failed to load any Whisper model!")
    return None, "mock-transcriber", False

# Try to load the best model available
model, MODEL_NAME, WHISPER_LOADED = load_best_whisper_model()

# Mock transcription for fallback
def mock_transcribe(audio_path: str) -> str:
    """Mock transcription function for testing"""
    return f"This is a mock transcription of the audio file: {audio_path}. The actual transcription would appear here once Whisper is properly installed and configured."

# Create directories
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Supported audio formats
SUPPORTED_FORMATS = {".mp3", ".wav", ".m4a", ".flac", ".ogg", ".wma", ".aac", ".mpeg", ".mpg", ".mp4"}

@app.get("/")
async def root():
    return {
        "message": "Record to Text API is running!", 
        "mode": "whisper" if WHISPER_LOADED else "mock",
        "model": MODEL_NAME,
        "environment": "development" if BACKEND_DEBUG else "production"
    }

@app.get("/health")
async def health_check():
    return {
        "status": "healthy", 
        "model": MODEL_NAME,
        "whisper_loaded": WHISPER_LOADED,
        "environment": "development" if BACKEND_DEBUG else "production",
        "timeout_limits": {
            "default": DEFAULT_TIMEOUT,
            "min": MIN_TIMEOUT,
            "max": MAX_TIMEOUT
        }
    }

def transcribe_audio(audio_path: str, language: Optional[str] = None, task: str = "transcribe") -> str:
    """Transcribe audio file using Whisper with advanced options"""
    try:
        if WHISPER_LOADED:
            # Advanced Whisper transcription with accuracy optimizations
            result = model.transcribe(
                audio_path,
                language=language,  # Auto-detect if None
                task=task,  # "transcribe" or "translate"
                verbose=False,
                word_timestamps=False,
                # Accuracy improvements
                beam_size=5,  # Higher beam size for better accuracy
                best_of=5,  # Generate multiple candidates and pick best
                temperature=0,  # Deterministic output for consistency
                condition_on_previous_text=True,  # Use context from previous segments
                compression_ratio_threshold=2.4,  # Skip low-quality segments
                logprob_threshold=-1.0,  # Skip uncertain segments
                no_speech_threshold=0.6,  # Better silence detection
                # Audio preprocessing
                fp16=torch.cuda.is_available(),  # Use FP16 if GPU available
            )
            return result["text"].strip()
        else:
            # Fallback to mock transcription
            return mock_transcribe(audio_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error transcribing audio: {str(e)}")

async def transcribe_audio_with_timeout(audio_path: str, language: Optional[str] = None, task: str = "transcribe", timeout_minutes: int = DEFAULT_TIMEOUT) -> str:
    """Transcribe audio with backend timeout handling"""
    def run_transcription():
        return transcribe_audio(audio_path, language, task)
    
    try:
        # Run transcription in thread pool with timeout
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future = executor.submit(run_transcription)
            try:
                # Wait for completion with timeout
                result = future.result(timeout=timeout_minutes * 60)  # Convert to seconds
                return result
            except concurrent.futures.TimeoutError:
                # Cancel the future and raise timeout error
                future.cancel()
                raise HTTPException(
                    status_code=408, 
                    detail=f"Transcription timed out after {timeout_minutes} minutes. Please try with a shorter audio file or contact support for longer files."
                )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error transcribing audio: {str(e)}")

def create_docx(text: str, filename: str) -> str:
    """Create a DOCX file with the transcribed text"""
    try:
        doc = Document()
        doc.add_heading('Audio Transcription', 0)
        doc.add_paragraph(text)
        
        output_path = OUTPUT_DIR / f"{filename}.docx"
        doc.save(output_path)
        return str(output_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating DOCX: {str(e)}")

def create_txt(text: str, filename: str) -> str:
    """Create a TXT file with the transcribed text"""
    try:
        output_path = OUTPUT_DIR / f"{filename}.txt"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Audio Transcription\n")
            f.write(f"==================\n")
            f.write(f"Source File: {filename}\n\n")
            f.write(text)
        return str(output_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating TXT: {str(e)}")

@app.post("/transcribe")
async def transcribe_audio_endpoint(
    file: UploadFile = File(...),
    output_format: str = Form("text"),
    language: str = Form(""),
    task: str = Form("transcribe"),
    timeout_minutes: int = Form(DEFAULT_TIMEOUT)  # Use environment default
):
    # Validate file
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    # Check file extension
    file_extension = Path(file.filename).suffix.lower()
    if file_extension not in SUPPORTED_FORMATS:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file format. Supported formats: {', '.join(SUPPORTED_FORMATS)}"
        )
    
    # Validate output format
    valid_output_formats = {"text", "docx", "txt"}
    if output_format not in valid_output_formats:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid output format. Supported formats: {', '.join(valid_output_formats)}"
        )
    
    # Validate timeout using environment limits
    if timeout_minutes < MIN_TIMEOUT or timeout_minutes > MAX_TIMEOUT:
        raise HTTPException(
            status_code=400, 
            detail=f"Timeout must be between {MIN_TIMEOUT} and {MAX_TIMEOUT} minutes"
        )
    
    try:
        # Create temporary directory for this request
        with tempfile.TemporaryDirectory() as temp_dir:
            # Save uploaded file
            upload_path = Path(temp_dir) / file.filename
            async with aiofiles.open(upload_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            
            # Prepare language parameter
            lang_param = language if language and language != "auto" else None
            
            # Transcribe audio with timeout
            transcribed_text = await transcribe_audio_with_timeout(
                str(upload_path), 
                lang_param, 
                task, 
                timeout_minutes
            )
            
            if output_format == "docx":
                # Create DOCX file
                docx_path = create_docx(transcribed_text, Path(file.filename).stem)
                
                # Return DOCX file
                return FileResponse(
                    docx_path,
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    filename=f"{Path(file.filename).stem}_transcription.docx"
                )
            elif output_format == "txt":
                # Create TXT file
                txt_path = create_txt(transcribed_text, Path(file.filename).stem)
                
                # Return TXT file
                return FileResponse(
                    txt_path,
                    media_type='text/plain',
                    filename=f"{Path(file.filename).stem}_transcription.txt"
                )
            else:
                # Return JSON response (text format)
                word_count = len(transcribed_text.split())
                character_count = len(transcribed_text)
                
                return {
                    "filename": file.filename,
                    "transcription": transcribed_text,
                    "word_count": word_count,
                    "character_count": character_count,
                    "processing_time_limit": f"{timeout_minutes} minutes",
                    "language": language if language else "auto-detected",
                    "task": task
                }
                
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

@app.get("/supported-formats")
async def get_supported_formats():
    """Get list of supported audio formats"""
    return {"supported_formats": list(SUPPORTED_FORMATS)}

if __name__ == "__main__":
    import uvicorn
    print(f"Starting server on {BACKEND_HOST}:{BACKEND_PORT}")
    uvicorn.run(app, host=BACKEND_HOST, port=BACKEND_PORT, log_level="debug" if BACKEND_DEBUG else "info") 